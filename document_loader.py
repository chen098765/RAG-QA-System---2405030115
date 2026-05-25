import os
from io import BytesIO
from PyPDF2 import PdfReader
from docx import Document

def read_pdf(file_input):
    """读取PDF文件内容，支持文件路径或文件对象"""
    text = ""
    try:
        # 判断是文件路径还是文件对象
        if isinstance(file_input, str):
            reader = PdfReader(file_input)
        else:
            reader = PdfReader(file_input)
        
        for page in reader.pages:
            text += page.extract_text() + "\n"
    except Exception as e:
        print(f"读取PDF文件失败: {e}")
    return text

def read_docx(file_input):
    """读取DOCX文件内容，支持文件路径或文件对象"""
    text = ""
    try:
        # 判断是文件路径还是文件对象
        if isinstance(file_input, str):
            doc = Document(file_input)
        else:
            doc = Document(BytesIO(file_input.read()))
        
        for para in doc.paragraphs:
            text += para.text + "\n"
    except Exception as e:
        print(f"读取DOCX文件失败: {e}")
    return text

def read_txt(file_input):
    """读取TXT文件内容，支持文件路径或文件对象"""
    text = ""
    try:
        # 判断是文件路径还是文件对象
        if isinstance(file_input, str):
            with open(file_input, 'r', encoding='utf-8') as f:
                text = f.read()
        else:
            text = file_input.read().decode('utf-8')
    except Exception as e:
        print(f"读取TXT文件失败: {e}")
    return text

def load_documents_from_folder(folder_path):
    """批量读取指定文件夹内的所有文档"""
    documents = []
    if not os.path.exists(folder_path):
        print(f"文件夹不存在: {folder_path}")
        return documents
    
    for filename in os.listdir(folder_path):
        file_path = os.path.join(folder_path, filename)
        if not os.path.isfile(file_path):
            continue
        
        _, ext = os.path.splitext(filename)
        text = ""
        
        if ext.lower() == '.pdf':
            text = read_pdf(file_path)
        elif ext.lower() == '.docx':
            text = read_docx(file_path)
        elif ext.lower() == '.txt':
            text = read_txt(file_path)
        else:
            print(f"不支持的文件格式: {filename}")
            continue
        
        if text.strip():
            documents.append({
                'filename': filename,
                'content': text
            })
            print(f"已读取: {filename}")
    
    return documents

if __name__ == "__main__":
    docs = load_documents_from_folder('docs')
    print(f"\n共读取 {len(docs)} 个文档")